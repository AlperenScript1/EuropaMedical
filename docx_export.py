import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Pt, RGBColor

FONT_ADI = "Times New Roman"
FONT_BOYUT = Pt(12)
TABLO_KOLON_GENISLIKLERI = [
    Emu(750000),   # Sıra No
    Emu(3600000),  # Cinsi (geniş)
    Emu(1100000),  # Miktarı
    Emu(1100000),  # Birim
]

HARIC_TUTULACAKLAR = [
    r"^Resmi dil\s*\(",
    r"^Pdf",
    r"^İmzalı pdf",
    r"^Makine çevirisi",
    r"^BG$|^CS$|^DA$|^DE$|^EL$|^ES$|^EN$|^ET$|^FI$|^FR$",
    r"^GA$|^HR$|^HU$|^IT$|^LT$|^LV$|^MT$|^NL$|^PL$|^PT$|^RO$|^SK$|^SL$|^SV$",
    r"^eTranslation",
    r"^Avrupa Komisyonu doğruluğu",
]

FORM_TURLERI = {
    "SONUÇ",
    "REKABET",
    "KONKUR",
    "İLAN",
    "COMPETITION",
    "RESULT",
    "CONTRACT",
    "CONTRACT NOTICE",
    "CONTRACT AWARD",
}

REKABET_FORM_TURLERI = {"REKABET", "KONKUR", "COMPETITION", "CONTRACT NOTICE", "İLAN"}

TAM_KALIN_TARIH_ETIKETLERI = [
    "Son teklif tarihi",
    "Son başvuru tarihi",
    "Tekliflerin son teslim tarihi",
    "İhalelerin son teslim tarihi",
    "İhalenin son teslim tarihi",
]

KISMI_KALIN_TARIH_ETIKETLERI = [
    "Sözleşmenin imzalanma tarihi",
    "Bildirimin gönderilme tarihi",
    "Yayınlanma tarihi",
    "Son tarih",
]

TARIH_KALIN_PARCA = re.compile(
    r"\d{2}/\d{2}/\d{4}|\d{2}:\d{2}(?::\d{2})?|\+\d{2}:\d{2}"
)


def _normalize_satir(satir: str) -> str:
    satir = re.sub(r"\s*:\s*", ":", satir.strip())
    satir = re.sub(r"\s+", " ", satir)
    return satir


def _haric_mi(satir: str) -> bool:
    if not satir.strip():
        return True
    if satir.strip() in {"=== ÖZET ===", "=== İLAN ==="}:
        return True
    if re.fullmatch(r"\d+\.\d*\.?", satir.strip()):
        return True
    for pattern in HARIC_TUTULACAKLAR:
        if re.search(pattern, satir.strip(), re.IGNORECASE):
            return True
    return False


def _alan_bul(pattern: str, metin: str, default: str = "") -> str:
    match = re.search(pattern, metin, re.IGNORECASE | re.MULTILINE)
    return match.group(1).strip() if match else default


def _tr_kucuk(metin: str) -> str:
    return metin.replace("I", "ı").replace("İ", "i").lower()


def _tr_baslik_kelime(kelime: str) -> str:
    kelime = _tr_kucuk(kelime)
    if not kelime:
        return kelime
    ilk = kelime[0]
    if ilk == "i":
        ilk = "İ"
    elif ilk == "ı":
        ilk = "I"
    else:
        ilk = ilk.upper()
    return ilk + kelime[1:]


def _bas_harf_buyuk(metin: str) -> str:
    return " ".join(_tr_baslik_kelime(kelime) for kelime in metin.split())


def _ulke_bul(summary: str, notice: str) -> str:
    for kaynak in (summary, notice):
        match = re.search(
            r"^([\wğüşıöçĞÜŞİÖÇ]+)\s*[:\–-]\s*",
            kaynak,
            re.MULTILINE | re.IGNORECASE,
        )
        if match:
            return match.group(1).strip().upper()
        match = re.search(r"Ülke\s*:\s*(.+)", kaynak, re.IGNORECASE)
        if match:
            return match.group(1).strip().upper()
    return "BİLİNMEYEN"


def _form_turu_bul(summary: str, notice: str) -> str:
    form = _alan_bul(r"Form türü\s*:\s*(.+)", notice)
    if form:
        return form.upper()
    ilk = summary.splitlines()[0].strip() if summary else ""
    if ilk:
        return ilk.upper()
    return "İLAN"


def _ihale_basligi_bul(summary: str, notice: str) -> str:
    metin = f"{summary}\n{notice}"
    lot = _alan_bul(r"LOT-0001\s*:\s*(.+)", metin)
    if lot:
        return _bas_harf_buyuk(lot)

    for ham in summary.splitlines():
        satir = _normalize_satir(ham)
        if not satir or _haric_mi(satir):
            continue
        if re.match(r"^[\wğüşıöçĞÜŞİÖÇ]+\s*[:\–-]\s*", satir):
            continue
        if re.fullmatch(r"\d+\.\s+[^:\.]+", satir):
            continue
        if len(satir) >= 15:
            return _bas_harf_buyuk(satir)

    return ""


def _alt_baslik_metni(form_turu: str, ihale_basligi: str) -> str:
    """Tablo/cinsi metni üstte tekrar gösterilmez; referans Word şablonunda alt başlık yok."""
    return ""


def _ozet_satirlari(
    summary: str, notice: str, form_turu: str, ihale_basligi: str
) -> list[str]:
    tablo_cinsi_ham = _tablo_cinsi_bul(summary, notice, ihale_basligi)
    tablo_cinsi, _, _ = _cinsi_miktar_ayir(tablo_cinsi_ham) if tablo_cinsi_ham else ("", "1", "Adet")

    satirlar = []
    for ham in summary.splitlines():
        satir = _normalize_satir(ham)
        if _haric_mi(satir):
            continue
        if satir.upper() == form_turu.upper():
            continue
        if ihale_basligi and satir.upper() == ihale_basligi.upper():
            continue
        if tablo_cinsi and _bas_harf_buyuk(satir).upper() == tablo_cinsi.upper():
            continue
        if tablo_cinsi_ham and _bas_harf_buyuk(satir).upper() == tablo_cinsi_ham.upper():
            continue
        if re.fullmatch(r"\d+\.\s+[^:\.]+", satir):
            continue
        satirlar.append(satir)
    return satirlar


def _org_ve_bildirim_satirlari(notice: str) -> list[str]:
    degisim_satirlari = []
    bildirim_satirlari = []
    org_satirlari = []
    mod = None

    for ham in notice.splitlines():
        satir = _normalize_satir(ham)
        if not satir:
            continue

        if satir.startswith("10. Change") or satir.startswith("10. Değiştirmek"):
            mod = "degisim"
            degisim_satirlari.append(satir)
            continue
        if satir.startswith("Bildirim bilgileri"):
            mod = "bildirim"
            bildirim_satirlari.append(satir)
            continue
        if satir.startswith("ORG-"):
            mod = "org"
            org_satirlari.append(satir)
            continue
        if re.fullmatch(r"8\.\s*Kuruluşlar", satir):
            mod = "org"
            continue

        if mod == "degisim" and not satir.startswith(("8.", "Bildirim")):
            if not _haric_mi(satir) and not re.fullmatch(r"\d+\.\s+[^:\.]+", satir):
                degisim_satirlari.append(satir)
            continue

        if mod == "org":
            if satir.startswith("Bildirim bilgileri"):
                mod = "bildirim"
                bildirim_satirlari.append(satir)
                continue
            if _haric_mi(satir) or re.fullmatch(r"8\.1\.", satir):
                continue
            org_satirlari.append(satir)
            continue

        if mod == "bildirim":
            if not _haric_mi(satir):
                bildirim_satirlari.append(satir)

    sonuc = []
    if org_satirlari:
        sonuc.append("Resmi dil (İmzalı PDF)")
        sonuc.extend(org_satirlari)
    if degisim_satirlari:
        sonuc.extend(degisim_satirlari)
    if bildirim_satirlari:
        sonuc.extend(bildirim_satirlari)
    return sonuc


def _ana_baslik_bul(summary: str) -> str:
    """Özet metnindeki ana büyük harfli ihale başlığını bulur."""
    for ham in summary.splitlines():
        satir = _normalize_satir(ham)
        if not satir or _haric_mi(satir):
            continue
        if re.match(r"^[\wğüşıöçĞÜŞİÖÇ]+\s*[:\–-]\s*", satir):
            continue
        if re.fullmatch(r"\d+\.\s+[^:\.]+", satir):
            continue
        if satir.startswith("LOT-"):
            continue

        harfler = [c for c in satir if c.isalpha()]
        if len(harfler) < 8:
            continue
        buyuk_oran = sum(1 for c in harfler if c.isupper()) / len(harfler)
        if buyuk_oran >= 0.75:
            return _bas_harf_buyuk(satir)
    return ""


def _lot_aciklamalari(summary: str, notice: str) -> list[str]:
    metin = f"{summary}\n{notice}"
    lotlar: dict[str, str] = {}
    for lot_no, aciklama in re.findall(r"(LOT-\d+)\s*:\s*(.+)", metin, re.IGNORECASE):
        lot_anahtar = lot_no.upper()
        if lot_anahtar in lotlar:
            continue
        aciklama = re.sub(r"^Lot\s+\d+\s*-\s*", "", aciklama.strip(), flags=re.IGNORECASE)
        if aciklama:
            lotlar[lot_anahtar] = aciklama
    return list(lotlar.values())


def _tablo_cinsi_bul(summary: str, notice: str, ihale_basligi: str) -> str:
    lotlar = _lot_aciklamalari(summary, notice)

    if len(lotlar) >= 2:
        ana_baslik = _ana_baslik_bul(summary)
        if ana_baslik:
            return ana_baslik

    if len(lotlar) == 1:
        return _bas_harf_buyuk(lotlar[0])

    if ihale_basligi:
        return ihale_basligi

    ana_baslik = _ana_baslik_bul(summary)
    if ana_baslik:
        return ana_baslik

    metin = f"{summary}\n{notice}"
    for _, aciklama in re.findall(
        r"(?:cpv\)|CPV\)|sınıflandırma\s*\(\s*cpv\s*\))\s*:\s*(\d{8})\s+(.+)",
        metin,
        re.IGNORECASE,
    ):
        return _bas_harf_buyuk(aciklama.strip())

    return ""


def _cinsi_miktar_ayir(cinsi: str) -> tuple[str, str, str]:
    """
    Cinsi sonundaki '- 11 Parça' / '11 Adet' gibi ifadeleri ayırır.
    Miktar ve birim tablo sütunlarına gider; cinsi sadece ürün/hizmet adını tutar.
    """
    for pattern in (
        r"\s*[-–—]\s*(\d+)\s*(?:Adet|Parça|Parca|Pcs|Pieces?|Units?)\.?\s*$",
        r"\s+(\d+)\s*(?:Adet|Parça|Parca|Pcs|Pieces?|Units?)\.?\s*$",
    ):
        match = re.search(pattern, cinsi, re.IGNORECASE)
        if match:
            temiz = cinsi[: match.start()].strip()
            temiz = re.sub(r"\s*[-–—]\s*$", "", temiz)
            return _bas_harf_buyuk(temiz), match.group(1), "Adet"

    return _bas_harf_buyuk(cinsi.strip()), "1", "Adet"


def _tablo_miktar_birim_bul(summary: str, notice: str) -> tuple[str, str]:
    metin = f"{summary}\n{notice}"

    for pattern in (
        r"(\d+)\s*(?:Adet|Parça|Parca)",
        r"(?:Tahmini miktar|Miktar|Estimated quantity)\s*:\s*(\d+)",
    ):
        match = re.search(pattern, metin, re.IGNORECASE)
        if match:
            return match.group(1), "Adet"

    return "1", "Adet"


def _tablo_satirlari(summary: str, notice: str, ihale_basligi: str) -> list[tuple[str, str, str, str]]:
    cinsi_ham = _tablo_cinsi_bul(summary, notice, ihale_basligi)
    if not cinsi_ham:
        return []

    cinsi, miktar, birim = _cinsi_miktar_ayir(cinsi_ham)
    if miktar == "1":
        govde_miktar, govde_birim = _tablo_miktar_birim_bul(summary, notice)
        if govde_miktar != "1":
            miktar, birim = govde_miktar, govde_birim

    return [("1", cinsi, miktar, birim)]


def _varsayilan_yazi_tipi_ayarla(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_ADI
    normal.font.size = FONT_BOYUT
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _run_ekle(paragraf, metin: str, kalin: bool = False):
    run = paragraf.add_run(metin)
    run.font.name = FONT_ADI
    run.font.size = FONT_BOYUT
    run.font.bold = kalin
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def _tarih_degeri_parcala(deger: str) -> list[tuple[str, bool]]:
    parcalar: list[tuple[str, bool]] = []
    son = 0
    for match in TARIH_KALIN_PARCA.finditer(deger):
        if match.start() > son:
            parcalar.append((deger[son : match.start()], False))
        parcalar.append((match.group(0), True))
        son = match.end()
    if son < len(deger):
        parcalar.append((deger[son:], False))
    return parcalar or [(deger, True)]


def _etiket_eslesir(etiket: str, adaylar: list[str]) -> bool:
    etiket = etiket.strip()
    return any(
        etiket.lower() == aday.lower() or etiket.lower().startswith(aday.lower())
        for aday in adaylar
    )


def _tam_kalin_etiket_mi(etiket: str) -> bool:
    return _etiket_eslesir(etiket, TAM_KALIN_TARIH_ETIKETLERI)


def _kismi_kalin_etiket_mi(etiket: str) -> bool:
    return _etiket_eslesir(etiket, KISMI_KALIN_TARIH_ETIKETLERI)


def _kalin_etiket_mi(etiket: str) -> bool:
    return _tam_kalin_etiket_mi(etiket) or _kismi_kalin_etiket_mi(etiket)


def _bos_satir_ekle(doc: Document):
    paragraf = doc.add_paragraph()
    paragraf.style = doc.styles["Normal"]
    paragraf.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _paragraf_ekle(doc: Document, satir: str, ortala: bool = False):
    paragraf = doc.add_paragraph()
    paragraf.style = doc.styles["Normal"]
    if ortala:
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraf.paragraph_format.space_after = Pt(6)
    else:
        paragraf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if ":" in satir:
        etiket, deger = satir.split(":", 1)
        etiket = etiket.strip()
        deger = deger.strip()
        _run_ekle(paragraf, f"{etiket}:", kalin=_tam_kalin_etiket_mi(etiket))
        if _tam_kalin_etiket_mi(etiket):
            _run_ekle(paragraf, deger, kalin=True)
        elif _kismi_kalin_etiket_mi(etiket):
            for parca, kalin in _tarih_degeri_parcala(deger):
                _run_ekle(paragraf, parca, kalin=kalin)
        else:
            _run_ekle(paragraf, deger)
    else:
        _run_ekle(paragraf, satir)

    return paragraf


def _hucre_yaz(hucre, metin: str, kalin: bool = False):
    hucre.text = ""
    hucre.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    paragraf = hucre.paragraphs[0]
    paragraf.paragraph_format.space_before = Pt(0)
    paragraf.paragraph_format.space_after = Pt(0)
    run = paragraf.add_run(metin)
    run.font.name = FONT_ADI
    run.font.size = FONT_BOYUT
    run.font.bold = kalin
    run.font.color.rgb = RGBColor(0, 0, 0)


def _tablo_ekle(doc: Document, satirlar: list[tuple[str, str, str, str]]):
    tablo = doc.add_table(rows=1 + len(satirlar), cols=4)
    tablo.style = "Table Grid"
    tablo.autofit = False

    basliklar = ["Sıra No", "Cinsi", "Miktarı", "Birim"]
    for ci, baslik in enumerate(basliklar):
        genislik = TABLO_KOLON_GENISLIKLERI[ci]
        tablo.columns[ci].width = genislik
        _hucre_yaz(tablo.rows[0].cells[ci], baslik)
        tablo.rows[0].cells[ci].width = genislik

    for row_idx, row_data in enumerate(satirlar, start=1):
        for col_idx, deger in enumerate(row_data):
            genislik = TABLO_KOLON_GENISLIKLERI[col_idx]
            hucre = tablo.rows[row_idx].cells[col_idx]
            hucre.width = genislik
            _hucre_yaz(hucre, deger)


def docx_olustur(
    summary: str,
    notice: str,
    ilan_no: str,
    url: str,
    cikti_yolu: Path,
) -> Path:
    ulke = _ulke_bul(summary, notice)
    form_turu = _form_turu_bul(summary, notice)
    ihale_basligi = _ihale_basligi_bul(summary, notice)
    alt_baslik = _alt_baslik_metni(form_turu, ihale_basligi)

    govde_satirlari = _ozet_satirlari(summary, notice, form_turu, ihale_basligi)
    govde_satirlari.extend(_org_ve_bildirim_satirlari(notice))
    if url:
        govde_satirlari.append(url)

    tablo_satirlari = _tablo_satirlari(summary, notice, ihale_basligi)

    doc = Document()
    _varsayilan_yazi_tipi_ayarla(doc)

    _paragraf_ekle(doc, f"YURT DIŞINDAN SATINALMA TALEBİ/ {ulke}", ortala=True)
    _bos_satir_ekle(doc)

    if alt_baslik:
        _paragraf_ekle(doc, alt_baslik, ortala=True)
        _bos_satir_ekle(doc)

    for satir in govde_satirlari:
        _paragraf_ekle(doc, satir)

    if tablo_satirlari:
        _tablo_ekle(doc, tablo_satirlari)

    doc.save(str(cikti_yolu))
    return cikti_yolu
